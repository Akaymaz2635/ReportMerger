from docx import Document
import pandas as pd
import re

def fill_inspection(template_path, output_path, data):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for match in re.findall(r'\[(.*?)\]', paragraph.text):  # Find all expressions within []
            if match in data:
                paragraph.text = paragraph.text.replace(f"[{match}]", str(data[match]))  # Replace placeholder with value

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for match in re.findall(r'\[(.*?)\]', paragraph.text):  # Find all expressions within []
                        if match in data:
                            paragraph.text = paragraph.text.replace(f"[{match}]", str(data[match]))  # Replace placeholder with value

    doc.save(output_path)

def generate_inspection_result_from_merged_values(template_path, excel_path):
    df = pd.read_excel(excel_path, sheet_name="TO_WORD")  # Read data from "TO_WORD" sheet
    for idx, row in df.iterrows():
        data = row.to_dict()  # Convert row to dictionary
        output_path = f'invitation_{idx + 1}.docx'
        fill_inspection(template_path, output_path, data)

if __name__ == '__main__':
    excel_path = "merged_data.xlsx"
    template_path = "template.docx"
    generate_inspection_result_from_merged_values(template_path, excel_path)
